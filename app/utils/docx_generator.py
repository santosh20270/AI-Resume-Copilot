from io import BytesIO
from docx import Document
from docx.shared import Pt


def generate_docx(markdown_text: str):

    doc = Document()

    title = doc.add_heading("AI Resume Copilot", level=1)
    title.runs[0].font.size = Pt(20)

    doc.add_heading("Rewritten Resume", level=2)

    lines = markdown_text.splitlines()

    for line in lines:

        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)

        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)

        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)

        elif line.startswith("- "):
            doc.add_paragraph(
                line.replace("- ", ""),
                style="List Bullet"
            )

        else:
            doc.add_paragraph(line)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer.getvalue()